import time
import sys
import rdflib
import PyPDF2
import docx
import re
import nltk
import subprocess
import json
import ols_client
import itertools
import datetime
import matplotlib.pyplot as plt
import pandas as pd

from Bio import Entrez
from gensim.models import doc2vec, word2vec
from collections import namedtuple
from sklearn.manifold import TSNE
from urllib.error import URLError
from SPARQLWrapper.SPARQLExceptions import EndPointInternalError


class OntoDoc:
    """Class to find ontology links based on a document, Pubmed ID or pasted text.
    """
    def __init__(self):
        """Initialisation
        """
        self.min_count = 1 # Minimum occurance of word
        self.min_length = 2 # Minimum word length
        self.epochs = 100 # Training iterations
        self.ols = ols_client.client.OlsClient() # Ontology Lookup Service client
        self.vector_size = 100 # Doc2Vec vector size
        self.window = 10 # Doc2Vec window size
        self.workers = 4 # Number of workers threads
        self.db = 'pubmed' # Entrez database to find pubmed abstracts
        self.retmode = 'xml' # Pubmed retrun mode
        self.regex = re.compile(r'([^\s\w]|_)+') # Regex to transform data before creating vectors
        self.stop_words = nltk.corpus.stopwords.words() # List of stop words
        self.link_percentage = 0.90 # Minimum word link percentage
        self.omicsscore = 20 # Minimum OMICS score on OmicsDI
        self.min_link = 4 # Minimum number of links in ontology description
        self.vocab = []  # The vocabulary from the Doc2Vec model
        self.tokens = [] # List with tokens from the Doc2Vec model
        self.tags = [] # List to store all available tags from data


    def pubmed_abstract(self, id_list):
        """Gets the abstract from an article in pubmed.
        The list with pubmed ids will be 

        Arguments:
            id_list: List of pubmed ids

        Returns:
            The abstracts from the pubmed article
        """
        abstracts = []
        Entrez.email = 'your.email@example.com'
        for id in id_list:
            abstract = ""
            handle = Entrez.efetch(db=self.db,
                                   retmode=self.retmode,
                                   id=id)
            results = Entrez.read(handle)
            for x in range(0, len(results['PubmedArticle'][0]['MedlineCitation']['Article']['Abstract']['AbstractText'])):
                abstract += results['PubmedArticle'][0]['MedlineCitation']['Article']['Abstract']['AbstractText'][x]
            abstracts.append(abstract)
        return abstracts


    def load_data(self, papers):
        """Loading the paper to read the content.

        Arguments:
            papers: Papers as a file location or as a pubmed abstract

        Returns:
            Sentences from the papers

        Raises:
            FileNotFoundError: There is no file found. 
        """
        doc = ""
        docs = {}
        sentences = []
        for i, paper in enumerate(papers):
            if ".pdf" in paper:
                pdfFileObject = open(paper, 'rb')
                pdfReader = PyPDF2.PdfFileReader(pdfFileObject)
                count = pdfReader.numPages
                for i in range(count):
                    page = pdfReader.getPage(i)
                    doc += page.extractText()
                sentences.append(self.regex.sub('', doc).lower().split("\n"))
            elif ".docx" in paper or ".doc" in paper:
                officedoc = docx.Document(paper)
                fullText = []
                for para in officedoc.paragraphs:
                    fullText.append(para.text)
                doc = '\n'.join(fullText)
                sentences.append(self.regex.sub('', doc).lower().split("\n"))
            else:
                try:
                    doc = open(paper, 'r').read()
                    doc = doc.replace('\n', '')
                    sentences.append(self.regex.sub('', doc).lower().split(". "))
                except FileNotFoundError:
                    docs[i] = paper.lower().split(". ")
        return docs


    def transform_data(self, docs):
        """Removes stop words.

        Arguments:
            sentences: A list with sentences from the paper
            min_length: The minimum word length

        Returns:
            List with analysed sentences
        """
        doclist = []
        analyzedDocument = namedtuple('AnalyzedDocument', 'words tags')
        for i, doc in docs.items():
            tags = [i]
            wordlist = []
            words = []
            for sentence in doc:
                splitsentence = sentence.split(" ")
                for s in splitsentence:
                    words.append(self.regex.sub('', s))
            for word in words:
                if word not in self.stop_words and len(word) >= self.min_length and word.isalpha():
                    wordlist.append(word)
            if wordlist:
                doclist.append(analyzedDocument(wordlist, tags))
        return doclist


    def train(self, doclist):
        """Trains the analysed document.

        Arguments:
            doc: List with the analysed paper
            min_count: Only words that occur more than min_count will be used
            epochs: The number of iterations over the data set 
            in order to train the model

        Returns:
            A doc2vec model
        """
        model = doc2vec.Doc2Vec(doclist,
                                vector_size=self.vector_size,
                                window=self.window,
                                min_count=self.min_count,
                                workers=self.workers)
        for dummyepoch in range(self.epochs):
            model.train(doclist, total_examples=model.corpus_count, epochs=1)
            model.alpha -= 0.0002
            model.min_alpha = model.alpha
        return model


    def plot(self, model):
        """Creates a scatterplot from the doc2vec model.

        Arguments:
            model: The doc2vec model

        Returns:
            A scatterplot with the doc2vec voabulary based on the submitted paper
            A list of words that can be used to tag data files connected to the paper
        """
        tags = []
        voc = list(model.wv.vocab)
        for v in voc:
            self.tokens.append(model[v])
            self.vocab.append(v)
        tsne = TSNE(perplexity=40, n_components=2,
                    init='pca', n_iter=2500, random_state=23)
        X_tsne = tsne.fit_transform(self.tokens)
        df = pd.DataFrame(X_tsne, index=self.vocab, columns=['x', 'y'])
        fig = plt.figure(figsize=(16, 16))
        ax = fig.add_subplot(1, 1, 1)
        ax.scatter(df['x'], df['y'], color='blue')
        for word, pos in df.iterrows():
            tags.append(word)
            ax.annotate(word,
                        pos)
        return tags


    def update_progress(self, job_title, progress):
        """Build a progress bar when searching databases.
        
        Arguments:
            job_title: Name of the job shown with the progressbar
            progress: Progress to calculate percentage
        """
        length = 100 # modify this to change the length
        block = int(round(length*progress))
        msg = "\r{0}: [{1}] {2}%".format(job_title, "#"*block + "-"*(length-block), round(progress*100, 2))
        if progress >= 1: msg += " DONE\r\n"
        sys.stdout.write(msg)
        sys.stdout.flush()


    def omics_di(self, tags, model):
        """Search the OmicsDI database for datasets related to tags found with Doc2Vec.
        
        Arguments:
            tags: Tags found by doc2vec
            model: doc2vec model
        
        Returns:
            A dictionary with all found OmicsDI URLs for each tag.
        """
        foundomics = {}
        tagnr = 0
        for tag in tags:
            tagnr += 1
            linked = model.wv.similar_by_word(tag)
            query = 'http://www.omicsdi.org/ws/dataset/search?query=' + tag
            res = subprocess.Popen(["curl", "-s", "-k", "-L", query], shell=True, stdout=subprocess.PIPE).communicate()[0].decode()
            jsonresult = json.loads(res)
            for x in range(0, len(jsonresult["datasets"])):
                count = 0
                description = jsonresult["datasets"][x]["description"]
                rid = jsonresult["datasets"][x]["id"]
                source = jsonresult["datasets"][x]["source"]
                score = int(float(jsonresult["datasets"][x]["viewsCountScaled"])*1000)
                url = ("https://www.omicsdi.org/dataset/" + source + "/" + rid)
                for link in linked:
                    if link[0] in description:
                        count += 1
                if count >= self.min_link and score >= self.omicsscore and link[1] >= self.link_percentage:
                    foundomics[rid] = url
                    # Test
            OntoDoc().update_progress("Searching tags in the OmicsDI database", tagnr/len(tags))
        return foundomics


    def ontologies(self, tags, model):
        """Search OLS to find ontologies that can be linked to the found tags.
        The words linked to the tag will be used to find the most appropriate 
        ontology by searching the ontology description.
        
        Arguments:
            tags: Tags found by doc2vec
            model: doc2vec model
        
        Returns:
            A dictionary with all found ontologies for each tag.
        
        Raises:
            KeyError: There is no ontology description
        """
        foundontologies = {}
        tagnr = 0
        for tag in tags:
            tagnr += 1
            linked = model.wv.similar_by_word(tag)
            ontolist = []
            searchonto = self.ols.search(tag)
            for i in range(0, len(searchonto['response']['docs'])):
                try:
                    iri = searchonto['response']['docs'][i]['iri']
                    label = searchonto['response']['docs'][i]['label'].lower()
                    description = searchonto['response']['docs'][i]['description'][0]
                    if iri not in ontolist:
                        if tag in label:
                            for link in linked:
                                if link[0] in description and link[1] > self.link_percentage:
                                    if iri not in ontolist:
                                        foundontologies[label] = iri
                                        ontolist.append(iri)
                except KeyError:
                    pass
            OntoDoc().update_progress("Searching tags in the OLS database", tagnr/len(tags))
        return foundontologies


    def disgenet(self, tags, model):
        """Send SPARQL query to DisGeNET endpoint to get URIs
        based on the provided tags.
        
        Arguments:
            tags: Found tags to use in DisGeNET search
            model: Doc2vec model to get the linked words for all tags.

        Returns:
            Dictionary with tgas and DisGeNET URIs.

        Raises:
            URLError, EndPointInternalError: SPARQL endpont is timed-out. 
            Can not find anything or server error.
        """
        disgenet_uris = {}
        tagnr = 0
        for tag in tags:
            try:
                tagnr += 1
                tag = tag.replace("'", "")
                linked = model.wv.similar_by_word(tag)
                sparql_query = (
                    "SELECT * " +
                    "WHERE { " +
                    "?disease rdf:type ncit:C7057 . " +
                    "?disease dcterms:title ?title . " +
                    "filter regex(?disease, \"umls/id\") . " +
                    "?title bif:contains \'\"" + tag + "\"\' ." +
                    "service <http://linkedlifedata.com/sparql> { " +
                    "?disease skos:definition ?description } " +
                    "}"
                )
                g = rdflib.ConjunctiveGraph('SPARQLStore')
                g.open("http://rdf.disgenet.org/sparql/")
                for row in g.query(sparql_query):
                    count = 0
                    for link in linked:
                        if link[0] in row[2] and link[1] > self.link_percentage:
                            count += 1
                    if count >= self.min_link:
                        disgenet_uris[row[1]] = row[0].strip("rdflib.term.URIRef")
                OntoDoc().update_progress("Searching tags in the DisGeNET triple store", tagnr/len(tags))
            except (URLError, EndPointInternalError):
                tagnr += 1
                OntoDoc().update_progress("Searching tags in the DisGeNET triple store", tagnr/len(tags))
                print("\n", tag, "not found!!!")
                pass
        return disgenet_uris


    def create_documents(self, found_data, tool):
        """Create OLS and DisGeNET document with all found ontology links.
        
        Arguments:
            foundontologies: Ontologies found with OLS and DisGeNET
        """
        with open(tool + ".txt", "w") as results:
            results.write("Data files found based on tags:\n\n")
            for name, iri in found_data.items():
                if iri:
                    results.write(name + ":\n")
                    results.write(iri + "\n")
                    results.write("\n")


    def onto_start(self, ontodoc):
        """Get all variables from user input.

            Returns:
                file input paths, minimum word occurance, minimum wword length, 
                number of training iterations and pubmed IDs
        """
        option = int(input("1 -- document; 2 -- Pubmed; 3 -- Paste text: "))
        while True:
            if option == 1:
                file_input = input("document paths (comma seperated): ")
                papers = file_input.split(',')
                break
            elif option == 2:
                ids = input("Enter pubmed ids (comma seperated): ")
                pubmed_ids = ids.split(',')
                papers = ontodoc.pubmed_abstract(pubmed_ids)
                break
            elif option == 3:
                papers = [input("Enter text to analyse: ")]
                break
            else:
                print("No valid option selected!")
                print("Please try again...")
                option = int(input("1 -- document; 2 -- Pubmed; 3 -- Paste text: "))
        toolset = input("o -- OmicsDI, d -- DisGeNET, l -- OLS: ")
        docs = ontodoc.load_data(papers)
        doclist = ontodoc.transform_data(docs)
        model = ontodoc.train(doclist)
        tags = ontodoc.plot(model)
        while toolset:
            if "o" in toolset:
                results = ontodoc.omics_di(tags, model)
                toolset = toolset.strip("o")
                tool = "OmicsDI"
            elif "d" in toolset:
                results = ontodoc.disgenet(tags, model)
                toolset = toolset.strip("d")
                tool = "DisGeNET"
            elif "l" in toolset:
                results = ontodoc.ontologies(tags, model)
                toolset = toolset.strip("l")
                tool = "OLS"
            else:
                print("No valid option!")
                print("Please try again")
                toolset = input("o -- OmicsDI, d -- DisGeNET, l -- OLS: ")
            ontodoc.create_documents(results, tool)
        plt.show()


if __name__ == "__main__":
    ontodoc = OntoDoc()
    ontodoc.onto_start(ontodoc)